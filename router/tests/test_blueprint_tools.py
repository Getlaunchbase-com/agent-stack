"""Acceptance tests for PR4: Blueprint takeoff toolchain.

Tests the full pipeline:
  PDF → extract_text → takeoff_low_voltage → xlsx + docx artifacts

Validates:
  - Output schema stability for all four tools
  - End-to-end pipeline produces JSON + XLSX + DOCX
  - All tools registered in /tools endpoint
  - Workspace validation (422) applies to blueprint tools
"""

import importlib
import json
import os
import zipfile

import fitz  # PyMuPDF — used to create sample PDFs
import openpyxl
import pytest
from docx import Document
from fastapi.testclient import TestClient


# ---- Fixtures ----

@pytest.fixture(autouse=True)
def _workspace_root(tmp_path, monkeypatch):
    """Set up a temp WORKSPACE_ROOT with a sample workspace + PDF."""
    ws_root = tmp_path / "workspaces"
    ws_root.mkdir()
    ws_dir = ws_root / "blueprint-proj"
    ws_dir.mkdir()

    monkeypatch.setenv("WORKSPACE_ROOT", str(ws_root))
    monkeypatch.setenv("ROUTER_AUTH_TOKEN", "")

    from router.app import workspace_tools, blueprint_tools, tools, main
    importlib.reload(workspace_tools)
    importlib.reload(blueprint_tools)
    importlib.reload(tools)
    importlib.reload(main)

    yield str(ws_root)


@pytest.fixture()
def ws_dir(_workspace_root):
    return os.path.join(_workspace_root, "blueprint-proj")


@pytest.fixture()
def client():
    from router.app.main import app
    return TestClient(app, raise_server_exceptions=False)


def _create_sample_pdf(ws_dir: str, filename: str = "plans/floor1.pdf") -> str:
    """Create a sample blueprint-like PDF with low-voltage device references."""
    full = os.path.join(ws_dir, filename)
    os.makedirs(os.path.dirname(full), exist_ok=True)

    doc = fitz.open()

    # Page 1: Data/network devices
    page1 = doc.new_page(width=792, height=612)  # 11x8.5 landscape
    page1.insert_text(
        (72, 72),
        "FIRST FLOOR LOW VOLTAGE PLAN\nDrawing: E-101\nScale: 1/8\" = 1'-0\"",
        fontsize=14,
    )
    page1.insert_text(
        (72, 150),
        "NETWORK CLOSET IDF-1A\n"
        "1x MDF Cabinet\n"
        "2x Patch Panel (48-port Cat6A)\n"
        "2x Network Switch (48-port PoE+)\n"
        "1x UPS (3kVA)\n",
        fontsize=10,
    )
    page1.insert_text(
        (72, 280),
        "OPEN OFFICE AREA\n"
        "24x Cat6A Data Drop\n"
        "6x Wireless AP (ceiling mount)\n"
        "4x CCTV IP Camera\n"
        "150 LF Conduit Run (3/4\" EMT)\n",
        fontsize=10,
    )

    # Page 2: Security/fire devices
    page2 = doc.new_page(width=792, height=612)
    page2.insert_text(
        (72, 72),
        "FIRST FLOOR SECURITY & FIRE ALARM PLAN\nDrawing: E-102",
        fontsize=14,
    )
    page2.insert_text(
        (72, 140),
        "MAIN ENTRANCE\n"
        "2x Card Reader (HID iClass SE)\n"
        "2x Door Contact\n"
        "2x REX Device\n"
        "1x Intercom Station\n",
        fontsize=10,
    )
    page2.insert_text(
        (72, 260),
        "FIRE ALARM DEVICES\n"
        "8x Smoke Detector (addressable)\n"
        "4x Horn/Strobe (wall mount)\n"
        "2x Pull Station\n"
        "3x Paging Speaker (70V)\n",
        fontsize=10,
    )

    doc.save(full)
    doc.close()
    return filename


@pytest.fixture()
def sample_pdf(ws_dir) -> str:
    """Create sample PDF and return its workspace-relative path."""
    return _create_sample_pdf(ws_dir)


# =====================================================================
# Test: Full pipeline — PDF → JSON → XLSX → DOCX
# =====================================================================

class TestFullPipeline:
    """Given a sample PDF in workspace, toolchain outputs JSON + XLSX + DOCX."""

    def test_end_to_end_pipeline(self, client, sample_pdf, ws_dir):
        """Full pipeline: extract → takeoff → xlsx + docx all produce valid artifacts."""
        # Step 1: Extract text from PDF
        r1 = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_extract_text",
                "arguments": {"workspace": "blueprint-proj", "pdf_path": sample_pdf},
            }
        })
        assert r1.status_code == 200
        extract = r1.json()
        assert extract["ok"] is True
        assert extract["page_count"] == 2
        pages = extract["pages"]

        # Step 2: Run takeoff on extracted pages
        r2 = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_takeoff_low_voltage",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "extracted_text": pages,
                    "project_name": "Test Building",
                    "drawing_number": "E-101/E-102",
                },
            }
        })
        assert r2.status_code == 200
        takeoff = r2.json()
        assert takeoff["ok"] is True

        # Step 3: Write XLSX
        r3 = client.post("/tool", json={
            "tool_call": {
                "name": "artifact_write_xlsx_takeoff",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "takeoff_json": takeoff,
                    "output_path": "artifacts/takeoff.xlsx",
                },
            }
        })
        assert r3.status_code == 200
        xlsx_result = r3.json()
        assert xlsx_result["ok"] is True
        assert os.path.isfile(xlsx_result["absolute_path"])

        # Step 4: Write DOCX
        r4 = client.post("/tool", json={
            "tool_call": {
                "name": "artifact_write_docx_summary",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "takeoff_json": takeoff,
                    "output_path": "artifacts/summary.docx",
                },
            }
        })
        assert r4.status_code == 200
        docx_result = r4.json()
        assert docx_result["ok"] is True
        assert os.path.isfile(docx_result["absolute_path"])

        # Verify all 3 artifacts on disk
        artifacts_dir = os.path.join(ws_dir, "artifacts")
        assert os.path.isfile(os.path.join(artifacts_dir, "takeoff.xlsx"))
        assert os.path.isfile(os.path.join(artifacts_dir, "summary.docx"))


# =====================================================================
# Test: Output schema stability — blueprint_extract_text
# =====================================================================

class TestExtractTextSchema:
    def test_schema_fields(self, client, sample_pdf):
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_extract_text",
                "arguments": {"workspace": "blueprint-proj", "pdf_path": sample_pdf},
            }
        })
        body = resp.json()
        assert body["ok"] is True
        # Top-level keys
        assert "pdf_path" in body
        assert "page_count" in body
        assert "pages" in body
        assert isinstance(body["pages"], list)

        # Per-page keys
        page = body["pages"][0]
        assert "page" in page
        assert isinstance(page["page"], int)
        assert "width" in page
        assert "height" in page
        assert "text" in page
        assert isinstance(page["text"], str)
        assert "blocks" in page
        assert isinstance(page["blocks"], list)

        # Block keys
        block = page["blocks"][0]
        assert "text" in block
        assert "bbox" in block
        assert isinstance(block["bbox"], list)
        assert len(block["bbox"]) == 4

    def test_missing_pdf_returns_error(self, client):
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_extract_text",
                "arguments": {"workspace": "blueprint-proj", "pdf_path": "nonexistent.pdf"},
            }
        })
        body = resp.json()
        assert body["ok"] is False
        assert "not found" in body["error"].lower()


# =====================================================================
# Test: Output schema stability — blueprint_takeoff_low_voltage
# =====================================================================

class TestTakeoffSchema:
    def test_schema_fields(self, client, sample_pdf):
        # Extract first
        r1 = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_extract_text",
                "arguments": {"workspace": "blueprint-proj", "pdf_path": sample_pdf},
            }
        })
        pages = r1.json()["pages"]

        # Takeoff
        r2 = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_takeoff_low_voltage",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "extracted_text": pages,
                    "project_name": "Schema Test",
                },
            }
        })
        body = r2.json()

        # Top-level schema
        assert body["ok"] is True
        assert "takeoff_id" in body
        assert "generated_at" in body
        assert "project_name" in body
        assert "summary" in body
        assert "line_items" in body
        assert "page_details" in body

        # Summary sub-schema
        summary = body["summary"]
        assert "total_device_types" in summary
        assert "total_devices" in summary
        assert isinstance(summary["total_device_types"], int)
        assert isinstance(summary["total_devices"], int)
        assert summary["total_device_types"] > 0
        assert summary["total_devices"] > 0

        # Line items sub-schema
        assert isinstance(body["line_items"], list)
        assert len(body["line_items"]) > 0
        item = body["line_items"][0]
        assert "device_type" in item
        assert "label" in item
        assert "unit" in item
        assert "quantity" in item
        assert isinstance(item["device_type"], str)
        assert isinstance(item["label"], str)
        assert isinstance(item["unit"], str)
        assert isinstance(item["quantity"], int)

        # Page details sub-schema
        assert isinstance(body["page_details"], list)
        assert len(body["page_details"]) > 0
        pd = body["page_details"][0]
        assert "page" in pd
        assert "devices" in pd
        assert isinstance(pd["devices"], dict)

    def test_detects_specific_devices(self, client, sample_pdf):
        """Verify the takeoff finds known devices from the sample PDF."""
        r1 = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_extract_text",
                "arguments": {"workspace": "blueprint-proj", "pdf_path": sample_pdf},
            }
        })
        pages = r1.json()["pages"]

        r2 = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_takeoff_low_voltage",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "extracted_text": pages,
                },
            }
        })
        body = r2.json()
        device_types = {item["device_type"] for item in body["line_items"]}

        # These should all be detected from our sample PDF
        assert "cat6_drop" in device_types
        assert "wireless_ap" in device_types
        assert "cctv_camera" in device_types
        assert "card_reader" in device_types
        assert "fire_alarm_device" in device_types
        assert "paging_speaker" in device_types

    def test_invalid_extracted_text_returns_error(self, client):
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_takeoff_low_voltage",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "extracted_text": "not a list",
                },
            }
        })
        body = resp.json()
        assert body["ok"] is False
        assert "list" in body["error"].lower()


# =====================================================================
# Test: Output schema stability — artifact_write_xlsx_takeoff
# =====================================================================

class TestXlsxSchema:
    def _get_takeoff(self, client, sample_pdf):
        r1 = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_extract_text",
                "arguments": {"workspace": "blueprint-proj", "pdf_path": sample_pdf},
            }
        })
        pages = r1.json()["pages"]
        r2 = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_takeoff_low_voltage",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "extracted_text": pages,
                    "project_name": "XLSX Test",
                },
            }
        })
        return r2.json()

    def test_schema_fields(self, client, sample_pdf):
        takeoff = self._get_takeoff(client, sample_pdf)
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "artifact_write_xlsx_takeoff",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "takeoff_json": takeoff,
                },
            }
        })
        body = resp.json()
        assert body["ok"] is True
        assert "artifact_path" in body
        assert "absolute_path" in body
        assert "size_bytes" in body
        assert "sheets" in body
        assert "row_count" in body
        assert isinstance(body["size_bytes"], int)
        assert body["size_bytes"] > 0
        assert "Takeoff Summary" in body["sheets"]

    def test_xlsx_is_valid_workbook(self, client, sample_pdf):
        takeoff = self._get_takeoff(client, sample_pdf)
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "artifact_write_xlsx_takeoff",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "takeoff_json": takeoff,
                },
            }
        })
        body = resp.json()
        # Load the XLSX and verify structure
        wb = openpyxl.load_workbook(body["absolute_path"])
        assert "Takeoff Summary" in wb.sheetnames
        ws = wb["Takeoff Summary"]
        # Header row at row 6
        assert ws.cell(row=6, column=1).value == "Device Type"
        assert ws.cell(row=6, column=4).value == "Quantity"
        # Data starts at row 7
        assert ws.cell(row=7, column=1).value is not None
        wb.close()

    def test_invalid_takeoff_json_returns_error(self, client):
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "artifact_write_xlsx_takeoff",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "takeoff_json": {"no_line_items": True},
                },
            }
        })
        body = resp.json()
        assert body["ok"] is False
        assert "line_items" in body["error"]


# =====================================================================
# Test: Output schema stability — artifact_write_docx_summary
# =====================================================================

class TestDocxSchema:
    def _get_takeoff(self, client, sample_pdf):
        r1 = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_extract_text",
                "arguments": {"workspace": "blueprint-proj", "pdf_path": sample_pdf},
            }
        })
        pages = r1.json()["pages"]
        r2 = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_takeoff_low_voltage",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "extracted_text": pages,
                    "project_name": "DOCX Test",
                },
            }
        })
        return r2.json()

    def test_schema_fields(self, client, sample_pdf):
        takeoff = self._get_takeoff(client, sample_pdf)
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "artifact_write_docx_summary",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "takeoff_json": takeoff,
                },
            }
        })
        body = resp.json()
        assert body["ok"] is True
        assert "artifact_path" in body
        assert "absolute_path" in body
        assert "size_bytes" in body
        assert "sections" in body
        assert isinstance(body["size_bytes"], int)
        assert body["size_bytes"] > 0
        assert "Summary" in body["sections"]
        assert "Device Schedule" in body["sections"]

    def test_docx_is_valid_document(self, client, sample_pdf):
        takeoff = self._get_takeoff(client, sample_pdf)
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "artifact_write_docx_summary",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "takeoff_json": takeoff,
                },
            }
        })
        body = resp.json()
        # Load DOCX and verify
        doc = Document(body["absolute_path"])
        # Should have title and headings (level=0 produces "Title" style)
        titled = [p.text for p in doc.paragraphs
                  if p.style.name.startswith("Heading") or p.style.name == "Title"]
        assert any("Takeoff" in h for h in titled)
        # Should have at least one table (device schedule)
        assert len(doc.tables) >= 1
        table = doc.tables[0]
        # Header row
        assert table.rows[0].cells[0].text == "Device Type"
        # Data rows
        assert len(table.rows) > 2  # header + at least 1 data + total

    def test_invalid_takeoff_json_returns_error(self, client):
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "artifact_write_docx_summary",
                "arguments": {
                    "workspace": "blueprint-proj",
                    "takeoff_json": "not a dict",
                },
            }
        })
        body = resp.json()
        assert body["ok"] is False
        assert "dict" in body["error"].lower()


# =====================================================================
# Test: Tools registered in /tools endpoint
# =====================================================================

class TestToolRegistration:
    def test_all_blueprint_tools_in_schema(self, client):
        resp = client.get("/tools")
        assert resp.status_code == 200
        names = [t["function"]["name"] for t in resp.json()["tools"]]
        assert "blueprint_extract_text" in names
        assert "blueprint_takeoff_low_voltage" in names
        assert "artifact_write_xlsx_takeoff" in names
        assert "artifact_write_docx_summary" in names

    def test_blueprint_tools_have_required_params(self, client):
        resp = client.get("/tools")
        tools_by_name = {
            t["function"]["name"]: t["function"] for t in resp.json()["tools"]
        }

        extract = tools_by_name["blueprint_extract_text"]
        assert "workspace" in extract["parameters"]["required"]
        assert "pdf_path" in extract["parameters"]["required"]

        takeoff = tools_by_name["blueprint_takeoff_low_voltage"]
        assert "workspace" in takeoff["parameters"]["required"]
        assert "extracted_text" in takeoff["parameters"]["required"]

        xlsx = tools_by_name["artifact_write_xlsx_takeoff"]
        assert "workspace" in xlsx["parameters"]["required"]
        assert "takeoff_json" in xlsx["parameters"]["required"]

        docx = tools_by_name["artifact_write_docx_summary"]
        assert "workspace" in docx["parameters"]["required"]
        assert "takeoff_json" in docx["parameters"]["required"]


# =====================================================================
# Test: Workspace validation applies to blueprint tools
# =====================================================================

class TestWorkspaceValidation:
    def test_extract_text_invalid_workspace_422(self, client):
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_extract_text",
                "arguments": {"workspace": "nonexistent", "pdf_path": "x.pdf"},
            }
        })
        assert resp.status_code == 422
        detail = resp.json()["detail"]
        assert detail["tool"] == "blueprint_extract_text"
        assert "availableWorkspaces" in detail

    def test_takeoff_invalid_workspace_422(self, client):
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "blueprint_takeoff_low_voltage",
                "arguments": {
                    "workspace": "nonexistent",
                    "extracted_text": [{"page": 1, "text": "test"}],
                },
            }
        })
        assert resp.status_code == 422

    def test_xlsx_invalid_workspace_422(self, client):
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "artifact_write_xlsx_takeoff",
                "arguments": {
                    "workspace": "nonexistent",
                    "takeoff_json": {"line_items": []},
                },
            }
        })
        assert resp.status_code == 422

    def test_docx_invalid_workspace_422(self, client):
        resp = client.post("/tool", json={
            "tool_call": {
                "name": "artifact_write_docx_summary",
                "arguments": {
                    "workspace": "nonexistent",
                    "takeoff_json": {"line_items": []},
                },
            }
        })
        assert resp.status_code == 422
