"""Blueprint parser toolchain (Tier-1).

Tools for processing construction blueprint PDFs into structured takeoff data
and generating XLSX/DOCX artifact outputs.

Pipeline:
  1. blueprint_extract_text   — PDF → extracted text + layout hints per page
  2. blueprint_takeoff_low_voltage — extracted text → structured JSON takeoff
  3. artifact_write_xlsx_takeoff  — takeoff JSON → XLSX spreadsheet artifact
  4. artifact_write_docx_summary  — takeoff JSON → DOCX summary artifact
"""

import json
import os
import re
import uuid
from datetime import datetime, timezone

import fitz  # PyMuPDF
import openpyxl
from docx import Document
from docx.shared import Inches, Pt
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter

WORKSPACE_ROOT = os.getenv("WORKSPACE_ROOT", "/workspaces")

# ---- Low-voltage device keyword catalogue ----
# Maps regex patterns to canonical device types for takeoff extraction.
_LV_DEVICE_PATTERNS: list[tuple[str, str]] = [
    (r"\bcat[\s\-]?6a?\b", "cat6_drop"),
    (r"\bdata\s+drop\b", "cat6_drop"),
    (r"\bnetwork\s+drop\b", "cat6_drop"),
    (r"\bfiber\s+drop\b", "fiber_drop"),
    (r"\bfiber\s+optic\b", "fiber_drop"),
    (r"\bwap\b", "wireless_ap"),
    (r"\bwireless\s+ap\b", "wireless_ap"),
    (r"\baccess\s+point\b", "wireless_ap"),
    (r"\bwi[\s\-]?fi\b", "wireless_ap"),
    (r"\bcctv\b", "cctv_camera"),
    (r"\bsecurity\s+camera\b", "cctv_camera"),
    (r"\bip\s+camera\b", "cctv_camera"),
    (r"\bcard\s+reader\b", "card_reader"),
    (r"\baccess\s+control\b", "card_reader"),
    (r"\bdoor\s+contact\b", "door_contact"),
    (r"\brex\b", "rex_device"),
    (r"\brequest[\s\-]to[\s\-]exit\b", "rex_device"),
    (r"\bfire\s+alarm\b", "fire_alarm_device"),
    (r"\bsmoke\s+detector\b", "fire_alarm_device"),
    (r"\bpull\s+station\b", "fire_alarm_device"),
    (r"\bhorn[\s/\-]?strobe\b", "fire_alarm_device"),
    (r"\bpaging\b", "paging_speaker"),
    (r"\bspeaker\b", "paging_speaker"),
    (r"\bintercom\b", "intercom"),
    (r"\bidf\b", "idf_cabinet"),
    (r"\bmdf\b", "mdf_cabinet"),
    (r"\bpatch\s+panel\b", "patch_panel"),
    (r"\bswitch\b", "network_switch"),
    (r"\bups\b", "ups"),
    (r"\bconduit\b", "conduit_run"),
    (r"\bj[\s\-]?hook\b", "j_hook"),
    (r"\bcable\s+tray\b", "cable_tray"),
]

# Device type metadata for reporting
_DEVICE_META: dict[str, dict] = {
    "cat6_drop": {"label": "Cat6/6A Data Drop", "unit": "ea"},
    "fiber_drop": {"label": "Fiber Optic Drop", "unit": "ea"},
    "wireless_ap": {"label": "Wireless Access Point", "unit": "ea"},
    "cctv_camera": {"label": "CCTV / IP Camera", "unit": "ea"},
    "card_reader": {"label": "Card Reader / Access Control", "unit": "ea"},
    "door_contact": {"label": "Door Contact", "unit": "ea"},
    "rex_device": {"label": "Request-to-Exit Device", "unit": "ea"},
    "fire_alarm_device": {"label": "Fire Alarm Device", "unit": "ea"},
    "paging_speaker": {"label": "Paging Speaker", "unit": "ea"},
    "intercom": {"label": "Intercom Station", "unit": "ea"},
    "idf_cabinet": {"label": "IDF Cabinet", "unit": "ea"},
    "mdf_cabinet": {"label": "MDF Cabinet", "unit": "ea"},
    "patch_panel": {"label": "Patch Panel", "unit": "ea"},
    "network_switch": {"label": "Network Switch", "unit": "ea"},
    "ups": {"label": "UPS", "unit": "ea"},
    "conduit_run": {"label": "Conduit Run", "unit": "lf"},
    "j_hook": {"label": "J-Hook", "unit": "ea"},
    "cable_tray": {"label": "Cable Tray", "unit": "lf"},
}


def _abs(workspace: str, path: str) -> str:
    """Resolve workspace-relative path with traversal protection."""
    if "/" in workspace or ".." in workspace:
        raise ValueError("Invalid workspace")
    base = os.path.join(WORKSPACE_ROOT, workspace)
    full = os.path.normpath(os.path.join(base, path))
    if not full.startswith(base):
        raise ValueError("Path traversal blocked")
    return full


# =====================================================================
# Tool 1: blueprint_extract_text
# =====================================================================

def blueprint_extract_text(workspace: str, pdf_path: str) -> dict:
    """Extract text and layout hints from a blueprint PDF.

    Returns per-page text blocks with bounding-box layout hints so
    downstream tools can reason about spatial placement on drawings.
    """
    full = _abs(workspace, pdf_path)
    if not os.path.isfile(full):
        return {"ok": False, "error": f"PDF not found: {pdf_path}"}

    try:
        doc = fitz.open(full)
    except Exception as e:
        return {"ok": False, "error": f"Failed to open PDF: {e}"}

    pages = []
    for page_num in range(len(doc)):
        page = doc[page_num]
        width = round(page.rect.width, 2)
        height = round(page.rect.height, 2)

        # Extract text blocks with positional data
        blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
        text_blocks = []
        full_text_parts = []

        for block in blocks:
            if block.get("type") != 0:  # skip image blocks
                continue
            bbox = block.get("bbox", (0, 0, 0, 0))
            block_text_parts = []
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    t = span.get("text", "").strip()
                    if t:
                        block_text_parts.append(t)
            block_text = " ".join(block_text_parts)
            if not block_text.strip():
                continue
            text_blocks.append({
                "text": block_text,
                "bbox": [round(v, 1) for v in bbox],
            })
            full_text_parts.append(block_text)

        pages.append({
            "page": page_num + 1,
            "width": width,
            "height": height,
            "text": "\n".join(full_text_parts),
            "blocks": text_blocks,
        })

    doc.close()

    return {
        "ok": True,
        "pdf_path": pdf_path,
        "page_count": len(pages),
        "pages": pages,
    }


# =====================================================================
# Tool 2: blueprint_takeoff_low_voltage
# =====================================================================

def _extract_quantity(text: str, match_start: int) -> int:
    """Try to find a numeric quantity near a device keyword match."""
    # Look backwards from match for a number like "12x" or "(12)" or just "12"
    prefix = text[max(0, match_start - 30):match_start]
    # Try patterns: "12x", "(12)", "12 ", "x12"
    qty_patterns = [
        r"(\d+)\s*[xX]\s*$",
        r"\((\d+)\)\s*$",
        r"(\d+)\s+$",
        r"[xX]\s*(\d+)\s*$",
        r"qty[:\s]*(\d+)",
    ]
    for pat in qty_patterns:
        m = re.search(pat, prefix, re.IGNORECASE)
        if m:
            return int(m.group(1))
    return 1


def _estimate_cable_length(text: str) -> int | None:
    """Try to extract linear footage from nearby text."""
    m = re.search(r"(\d+)\s*(?:lf|linear\s*f(?:ee)?t|ft)", text, re.IGNORECASE)
    if m:
        return int(m.group(1))
    return None


def blueprint_takeoff_low_voltage(
    workspace: str,
    extracted_text: list[dict],
    project_name: str = "",
    drawing_number: str = "",
) -> dict:
    """Produce a structured low-voltage takeoff from extracted blueprint text.

    Args:
        workspace: workspace ID
        extracted_text: list of page dicts from blueprint_extract_text
            (each with "page", "text", and optionally "blocks")
        project_name: optional project name for the report header
        drawing_number: optional drawing/sheet number
    Returns:
        Structured JSON with device counts, cable lengths, and per-page breakdown.
    """
    # Validate input structure
    if not isinstance(extracted_text, list):
        return {"ok": False, "error": "extracted_text must be a list of page dicts"}

    takeoff_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()

    # Aggregate counts across all pages
    device_totals: dict[str, int] = {}
    page_details: list[dict] = []

    for page_data in extracted_text:
        if not isinstance(page_data, dict):
            continue
        page_num = page_data.get("page", 0)
        text = page_data.get("text", "")
        if not text:
            continue

        text_lower = text.lower()
        page_devices: dict[str, int] = {}

        for pattern, device_type in _LV_DEVICE_PATTERNS:
            for m in re.finditer(pattern, text_lower):
                qty = _extract_quantity(text_lower, m.start())
                page_devices[device_type] = page_devices.get(device_type, 0) + qty

        # Cable length estimation
        cable_lf = _estimate_cable_length(text)

        if page_devices or cable_lf:
            detail = {"page": page_num, "devices": page_devices}
            if cable_lf:
                detail["estimated_cable_lf"] = cable_lf
            page_details.append(detail)

        for dev, qty in page_devices.items():
            device_totals[dev] = device_totals.get(dev, 0) + qty

    # Build line items with metadata
    line_items = []
    for device_type, count in sorted(device_totals.items()):
        meta = _DEVICE_META.get(device_type, {"label": device_type, "unit": "ea"})
        line_items.append({
            "device_type": device_type,
            "label": meta["label"],
            "unit": meta["unit"],
            "quantity": count,
        })

    takeoff = {
        "ok": True,
        "takeoff_id": takeoff_id,
        "generated_at": now,
        "project_name": project_name or "Untitled Project",
        "drawing_number": drawing_number,
        "summary": {
            "total_device_types": len(line_items),
            "total_devices": sum(item["quantity"] for item in line_items),
        },
        "line_items": line_items,
        "page_details": page_details,
    }
    return takeoff


# =====================================================================
# Tool 3: artifact_write_xlsx_takeoff
# =====================================================================

def artifact_write_xlsx_takeoff(
    workspace: str,
    takeoff_json: dict,
    output_path: str = "artifacts/takeoff.xlsx",
) -> dict:
    """Write a structured takeoff JSON to an XLSX spreadsheet artifact.

    Args:
        workspace: workspace ID
        takeoff_json: takeoff dict from blueprint_takeoff_low_voltage
        output_path: relative path in workspace for the output file
    """
    full = _abs(workspace, output_path)
    os.makedirs(os.path.dirname(full), exist_ok=True)

    # Validate required fields
    if not isinstance(takeoff_json, dict):
        return {"ok": False, "error": "takeoff_json must be a dict"}
    line_items = takeoff_json.get("line_items")
    if not isinstance(line_items, list):
        return {"ok": False, "error": "takeoff_json.line_items must be a list"}

    project_name = takeoff_json.get("project_name", "Untitled")
    drawing_number = takeoff_json.get("drawing_number", "")
    generated_at = takeoff_json.get("generated_at", "")
    takeoff_id = takeoff_json.get("takeoff_id", "")

    wb = openpyxl.Workbook()

    # ---- Sheet 1: Summary ----
    ws = wb.active
    ws.title = "Takeoff Summary"

    header_font = Font(bold=True, size=14)
    subheader_font = Font(bold=True, size=11)
    header_fill = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")
    header_text = Font(bold=True, color="FFFFFF", size=11)

    ws.merge_cells("A1:D1")
    ws["A1"] = f"Low-Voltage Takeoff — {project_name}"
    ws["A1"].font = header_font

    ws["A2"] = "Drawing:"
    ws["A2"].font = subheader_font
    ws["B2"] = drawing_number
    ws["A3"] = "Generated:"
    ws["A3"].font = subheader_font
    ws["B3"] = generated_at
    ws["A4"] = "Takeoff ID:"
    ws["A4"].font = subheader_font
    ws["B4"] = takeoff_id

    # Column headers
    headers = ["Device Type", "Description", "Unit", "Quantity"]
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=6, column=col, value=h)
        cell.font = header_text
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    # Data rows
    for i, item in enumerate(line_items):
        row = 7 + i
        ws.cell(row=row, column=1, value=item.get("device_type", ""))
        ws.cell(row=row, column=2, value=item.get("label", ""))
        ws.cell(row=row, column=3, value=item.get("unit", "ea"))
        ws.cell(row=row, column=4, value=item.get("quantity", 0))
        ws.cell(row=row, column=4).alignment = Alignment(horizontal="center")

    # Totals row
    total_row = 7 + len(line_items)
    ws.cell(row=total_row, column=1, value="TOTAL").font = Font(bold=True)
    ws.cell(row=total_row, column=4, value=sum(
        item.get("quantity", 0) for item in line_items
    ))
    ws.cell(row=total_row, column=4).font = Font(bold=True)
    ws.cell(row=total_row, column=4).alignment = Alignment(horizontal="center")

    # Column widths
    ws.column_dimensions["A"].width = 22
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 8
    ws.column_dimensions["D"].width = 12

    # ---- Sheet 2: Page Details ----
    page_details = takeoff_json.get("page_details", [])
    if page_details:
        ws2 = wb.create_sheet("Page Details")
        detail_headers = ["Page", "Device Type", "Description", "Quantity", "Est. Cable (LF)"]
        for col, h in enumerate(detail_headers, 1):
            cell = ws2.cell(row=1, column=col, value=h)
            cell.font = header_text
            cell.fill = header_fill
            cell.alignment = Alignment(horizontal="center")

        row = 2
        for pd in page_details:
            page_num = pd.get("page", "?")
            cable_lf = pd.get("estimated_cable_lf", "")
            for dev_type, qty in pd.get("devices", {}).items():
                meta = _DEVICE_META.get(dev_type, {"label": dev_type})
                ws2.cell(row=row, column=1, value=page_num)
                ws2.cell(row=row, column=2, value=dev_type)
                ws2.cell(row=row, column=3, value=meta.get("label", dev_type))
                ws2.cell(row=row, column=4, value=qty)
                ws2.cell(row=row, column=5, value=cable_lf if cable_lf else "")
                row += 1

        ws2.column_dimensions["A"].width = 8
        ws2.column_dimensions["B"].width = 22
        ws2.column_dimensions["C"].width = 30
        ws2.column_dimensions["D"].width = 12
        ws2.column_dimensions["E"].width = 16

    wb.save(full)
    file_size = os.path.getsize(full)

    return {
        "ok": True,
        "artifact_path": output_path,
        "absolute_path": full,
        "size_bytes": file_size,
        "sheets": ["Takeoff Summary"] + (["Page Details"] if page_details else []),
        "row_count": len(line_items),
    }


# =====================================================================
# Tool 4: artifact_write_docx_summary
# =====================================================================

def artifact_write_docx_summary(
    workspace: str,
    takeoff_json: dict,
    output_path: str = "artifacts/takeoff_summary.docx",
) -> dict:
    """Write a DOCX summary document from a structured takeoff JSON.

    Args:
        workspace: workspace ID
        takeoff_json: takeoff dict from blueprint_takeoff_low_voltage
        output_path: relative path in workspace for the output file
    """
    full = _abs(workspace, output_path)
    os.makedirs(os.path.dirname(full), exist_ok=True)

    if not isinstance(takeoff_json, dict):
        return {"ok": False, "error": "takeoff_json must be a dict"}
    line_items = takeoff_json.get("line_items")
    if not isinstance(line_items, list):
        return {"ok": False, "error": "takeoff_json.line_items must be a list"}

    project_name = takeoff_json.get("project_name", "Untitled")
    drawing_number = takeoff_json.get("drawing_number", "")
    generated_at = takeoff_json.get("generated_at", "")
    takeoff_id = takeoff_json.get("takeoff_id", "")
    summary = takeoff_json.get("summary", {})

    doc = Document()

    # Title
    doc.add_heading(f"Low-Voltage Takeoff Report", level=0)

    # Metadata paragraph
    meta_para = doc.add_paragraph()
    meta_para.add_run("Project: ").bold = True
    meta_para.add_run(project_name + "\n")
    if drawing_number:
        meta_para.add_run("Drawing: ").bold = True
        meta_para.add_run(drawing_number + "\n")
    meta_para.add_run("Generated: ").bold = True
    meta_para.add_run(generated_at + "\n")
    meta_para.add_run("Takeoff ID: ").bold = True
    meta_para.add_run(takeoff_id)

    # Summary section
    doc.add_heading("Summary", level=1)
    total_types = summary.get("total_device_types", len(line_items))
    total_devices = summary.get("total_devices", sum(i.get("quantity", 0) for i in line_items))
    doc.add_paragraph(
        f"This takeoff identified {total_types} distinct device type(s) "
        f"totalling {total_devices} device(s)."
    )

    # Device table
    doc.add_heading("Device Schedule", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Device Type"
    hdr[1].text = "Description"
    hdr[2].text = "Unit"
    hdr[3].text = "Qty"
    for cell in hdr:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    for item in line_items:
        row = table.add_row().cells
        row[0].text = item.get("device_type", "")
        row[1].text = item.get("label", "")
        row[2].text = item.get("unit", "ea")
        row[3].text = str(item.get("quantity", 0))

    # Totals row
    total_row = table.add_row().cells
    total_row[0].text = "TOTAL"
    total_row[1].text = ""
    total_row[2].text = ""
    total_row[3].text = str(total_devices)
    for cell in total_row:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True

    # Page details
    page_details = takeoff_json.get("page_details", [])
    if page_details:
        doc.add_heading("Page-by-Page Breakdown", level=1)
        for pd in page_details:
            page_num = pd.get("page", "?")
            doc.add_heading(f"Page {page_num}", level=2)
            devices = pd.get("devices", {})
            if devices:
                for dev_type, qty in devices.items():
                    meta = _DEVICE_META.get(dev_type, {"label": dev_type})
                    doc.add_paragraph(
                        f"{meta.get('label', dev_type)}: {qty}",
                        style="List Bullet",
                    )
            cable_lf = pd.get("estimated_cable_lf")
            if cable_lf:
                doc.add_paragraph(f"Estimated cable: {cable_lf} LF", style="List Bullet")

    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.add_run("Generated by LaunchBase Agent Stack — Blueprint Takeoff Toolchain").italic = True

    doc.save(full)
    file_size = os.path.getsize(full)

    return {
        "ok": True,
        "artifact_path": output_path,
        "absolute_path": full,
        "size_bytes": file_size,
        "sections": ["Summary", "Device Schedule"]
        + (["Page-by-Page Breakdown"] if page_details else []),
    }
